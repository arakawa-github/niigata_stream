from docx import Document
import pandas as pd
import re
import streamlit as st
import io

st.title("ファイルアップロード")

# ファイルアップローダー
uploaded_file = st.file_uploader("ファイルを選択してください", type=["docx"])

if uploaded_file is not None:
    # `uploaded_file` を `io.BytesIO` に変換
    file_stream = io.BytesIO(uploaded_file.read())

    # Wordファイルを開く
    doc = Document(file_stream)

    # データを格納するリスト
    data = []
    current_serial_number = None  # 連番を保持する変数
    tables = iter(doc.tables)  # すべての表を順番に処理するイテレータ

    # Word内のテキストを解析
    for para in doc.paragraphs:
        text = para.text.strip()

        # 「連番 H318」のようなフォーマットから H318 部分を抽出
        match = re.search(r"連番\s*([A-Za-z0-9]+)", text)
        if match:
            current_serial_number = match.group(1)  # 連番を更新
            
            # 次の表を取得（連番がある段落の後に表がある前提）
            try:
                table = next(tables)  # 次の表を取得
                table_data = []

                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    
                    # データが空白行でない場合のみ処理
                    if any(row_data):
                        table_data.append(row_data)

                # **表の最初の行（列名）を削除**
                if table_data:
                    table_data.pop(0)  

                # **処理後の表データを追加**
                for row in table_data:
                    if current_serial_number:  # 連番が取得できている場合のみ
                        data.append([current_serial_number] + row)

            except StopIteration:
                pass  # 表がもうない場合はスキップ

    # DataFrameを作成（最初の列は「連番」、残りは表のデータ）
    df = pd.DataFrame(data)

    # 列数が足りない場合のチェック
    if df.shape[1] >= 11:
        # 4列目（インデックス3）、9列目（インデックス8）、10列目（インデックス9）、11列目（インデックス10）を削除
        df = df.drop(df.columns[[3, 7, 8, 9, 10]], axis=1)

    # 列名を設定
    df.columns = ["連番", "棚番", "品目コード", "指示数", "納入日", "品名"]
    
    # 新しい列を追加（空の値で初期化）
    df["済み"] = ""
    df["担当者"] = ""
    df["ピッキング日"] = ""
    df["ID_XX"] = ""

    # ExcelファイルをBytesIOに保存
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine="openpyxl") as writer:
        df.to_excel(writer, index=False)
    output.seek(0)

    # ダウンロードボタン
    st.download_button(
        label="Excelをダウンロード",
        data=output,
        file_name="ピッキングリスト.xlsx",
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )
